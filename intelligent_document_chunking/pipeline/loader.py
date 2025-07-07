import os
import fitz  # PyMuPDF
from docx import Document
import markdown
from typing import List, Dict, Optional, Union
from pathlib import Path
import re
from dataclasses import dataclass

@dataclass
class DocumentMetadata:
    """Metadata for loaded documents."""
    filename: str
    file_path: str
    file_type: str
    file_size: int
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    created_date: Optional[str] = None
    modified_date: Optional[str] = None

class DocumentLoader:
    """Handles loading and preprocessing of various document formats."""
    
    def __init__(self):
        self.supported_extensions = {
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.doc': self._load_docx,
            '.md': self._load_markdown,
            '.txt': self._load_text,
            '.html': self._load_html
        }
    
    def load_document(self, file_path: str) -> Dict[str, Union[str, DocumentMetadata]]:
        """
        Load a document and return its content and metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with 'content' and 'metadata' keys
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Get file metadata
        metadata = self._get_file_metadata(file_path)
        
        # Load content
        content = self.supported_extensions[file_extension](file_path)
        
        # Update metadata with content info
        metadata.word_count = len(content.split())
        
        return {
            'content': content,
            'metadata': metadata
        }
    
    def _get_file_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract basic file metadata."""
        stat = file_path.stat()
        
        return DocumentMetadata(
            filename=file_path.name,
            file_path=str(file_path),
            file_type=file_path.suffix.lower(),
            file_size=stat.st_size,
            created_date=str(stat.st_ctime),
            modified_date=str(stat.st_mtime)
        )
    
    def _load_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            doc = fitz.open(str(file_path))
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
                
                # Add page separator for better chunking
                if page_num < len(doc) - 1:
                    text += "\n\n--- PAGE BREAK ---\n\n"
            
            doc.close()
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"Error loading PDF {file_path}: {e}")
    
    def _load_docx(self, file_path: Path) -> str:
        """Extract text from Word document."""
        try:
            doc = Document(str(file_path))
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                text += "\n--- TABLE ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
                text += "--- END TABLE ---\n\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"Error loading DOCX {file_path}: {e}")
    
    def _load_markdown(self, file_path: Path) -> str:
        """Load and parse markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Convert markdown to plain text while preserving structure
            text = self._markdown_to_text(content)
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"Error loading Markdown {file_path}: {e}")
    
    def _load_text(self, file_path: Path) -> str:
        """Load plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return self._clean_text(content)
            
        except Exception as e:
            raise Exception(f"Error loading text file {file_path}: {e}")
    
    def _load_html(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Simple HTML tag removal
            text = re.sub(r'<[^>]+>', '', content)
            text = re.sub(r'\s+', ' ', text)
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"Error loading HTML {file_path}: {e}")
    
    def _markdown_to_text(self, markdown_content: str) -> str:
        """Convert markdown to plain text while preserving structure."""
        # Preserve headers
        markdown_content = re.sub(r'^### (.*$)', r'HEADER3: \1', markdown_content, flags=re.MULTILINE)
        markdown_content = re.sub(r'^## (.*$)', r'HEADER2: \1', markdown_content, flags=re.MULTILINE)
        markdown_content = re.sub(r'^# (.*$)', r'HEADER1: \1', markdown_content, flags=re.MULTILINE)
        
        # Preserve code blocks
        markdown_content = re.sub(r'```(\w+)?\n(.*?)```', r'CODE_BLOCK:\2', markdown_content, flags=re.DOTALL)
        markdown_content = re.sub(r'`([^`]+)`', r'INLINE_CODE:\1', markdown_content)
        
        # Preserve lists
        markdown_content = re.sub(r'^\* (.*$)', r'BULLET: \1', markdown_content, flags=re.MULTILINE)
        markdown_content = re.sub(r'^\d+\. (.*$)', r'NUMBERED: \1', markdown_content, flags=re.MULTILINE)
        
        # Preserve bold and italic
        markdown_content = re.sub(r'\*\*(.*?)\*\*', r'BOLD:\1', markdown_content)
        markdown_content = re.sub(r'\*(.*?)\*', r'ITALIC:\1', markdown_content)
        
        # Remove other markdown syntax
        markdown_content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', markdown_content)  # Links
        markdown_content = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'IMAGE: \1', markdown_content)  # Images
        
        return markdown_content
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but preserve important ones
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\(\)\[\]\{\}\-\–\—\n\t]', '', text)
        
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        # Normalize dashes
        text = text.replace('–', '-').replace('—', '-')
        
        return text.strip()
    
    def load_multiple_documents(self, file_paths: List[str]) -> List[Dict[str, Union[str, DocumentMetadata]]]:
        """
        Load multiple documents.
        
        Args:
            file_paths: List of file paths to load
            
        Returns:
            List of document dictionaries
        """
        documents = []
        
        for file_path in file_paths:
            try:
                doc = self.load_document(file_path)
                documents.append(doc)
            except Exception as e:
                print(f"Warning: Could not load {file_path}: {e}")
                continue
        
        return documents
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self.supported_extensions.keys())
    
    def validate_file(self, file_path: str) -> bool:
        """Validate if a file can be loaded."""
        try:
            file_path = Path(file_path)
            return file_path.suffix.lower() in self.supported_extensions
        except:
            return False 